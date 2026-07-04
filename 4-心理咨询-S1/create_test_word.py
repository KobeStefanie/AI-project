#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建一个包含颜色的测试Word文档
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

# 创建文档
doc = Document()

# 添加段落1：红色文字
p1 = doc.add_paragraph()
run1 = p1.add_run('案例概要：')
run1.bold = True
run1.font.color.rgb = RGBColor(255, 0, 0)  # 红色
run2 = p1.add_run('14岁女孩的同伴来电求助')
run2.font.color.rgb = RGBColor(255, 0, 0)  # 红色
# 尝试添加背景色（通过高亮）
run2.font.highlight_color = WD_COLOR_INDEX.RED

# 添加段落2：蓝色文字
p2 = doc.add_paragraph()
run3 = p2.add_run('相关标签：')
run3.bold = True
run3.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

run4 = p2.add_run('社交关系-友谊-危机干预')
run4.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

# 添加段落3：黄色高亮
p3 = doc.add_paragraph()
run5 = p3.add_run('危机评估：')
run5.bold = True
run6 = p3.add_run('证据')
run6.font.highlight_color = WD_COLOR_INDEX.YELLOW

# 保存
output_path = 'test_color.docx'
doc.save(output_path)
print(f'测试文档已创建: {output_path}')
print('\n这个文档包含：')
print('- 红色文字："案例概要："和"14岁女孩的同伴来电求助"')
print('- 红色高亮背景："14岁女孩的同伴来电求助"')
print('- 蓝色文字："相关标签："和内容')
print('- 黄色高亮："证据"')
