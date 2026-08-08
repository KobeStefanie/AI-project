#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
下载文件生成脚本
功能：为每个案例生成可下载的文件（Word复盘、Excel逐字稿、流派分析Markdown）
"""

import os
import sys
import json
from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill

# ---- Windows GBK 兼容处理 ----
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def sp(*args, **kwargs):
    """安全 print：自动处理编码问题"""
    try:
        print(*args, **kwargs)
    except UnicodeEncodeError:
        safe_args = []
        for a in args:
            s = str(a)
            safe_args.append(s.encode('utf-8', errors='replace').decode('utf-8', errors='replace'))
        print(*safe_args, **kwargs)


# ==================== 配置 ====================

PROJECT_ROOT = Path(__file__).parent.parent
DATA_DIR = PROJECT_ROOT / "data"
CASES_PROCESSED = DATA_DIR / "cases" / "processed"
OUTPUT_DIR = PROJECT_ROOT / "output" / "案例库" / "downloads"


# ==================== Word生成：复盘 ====================

def generate_review_word(case_data: Dict, output_path: Path):
    """生成复盘Word文档"""
    doc = Document()

    # 标题
    title = doc.add_heading(f"咨询师复盘 - {case_data.get('case_id', 'Unknown')}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 案例基本信息
    doc.add_heading('基本信息', level=2)
    basic_info = case_data.get('basic_info', {})
    p = doc.add_paragraph()
    p.add_run(f"案例编号：{case_data.get('case_id', 'N/A')}\n")
    p.add_run(f"来访时间：{basic_info.get('date', 'N/A')}\n")
    p.add_run(f"性别：{basic_info.get('gender', 'N/A')}\n")
    p.add_run(f"年龄段：{basic_info.get('age_group', 'N/A')}\n")

    # 案例概要
    doc.add_heading('案例概要', level=2)
    summary = case_data.get('case_summary', '暂无概要')
    doc.add_paragraph(summary)

    # 完整对话记录
    doc.add_heading('完整对话记录', level=2)
    dialogue = case_data.get('dialogue', '暂无对话记录')
    doc.add_paragraph(dialogue)

    # 保存
    doc.save(output_path)
    sp(f"[OK] Word生成: {output_path.name}")


# ==================== Excel生成：逐字稿 ====================

def generate_transcript_excel(case_data: Dict, output_path: Path):
    """生成逐字稿Excel文档"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "逐字稿"

    # 设置列宽
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 80

    # 标题行
    headers = ['时间', '角色', '内容']
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True, size=12)
        cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        cell.font = Font(bold=True, size=12, color="FFFFFF")
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # 获取逐字稿数据
    transcripts = case_data.get('transcripts', [])

    if not transcripts:
        # 如果没有逐字稿，添加提示信息
        ws.cell(row=2, column=1, value='暂无逐字稿数据')
        ws.merge_cells('A2:C2')
        cell = ws.cell(row=2, column=1)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.font = Font(italic=True, color="999999")
    else:
        # 填充数据
        for idx, item in enumerate(transcripts, 2):
            ws.cell(row=idx, column=1, value=item.get('timestamp', ''))
            ws.cell(row=idx, column=2, value=item.get('speaker', ''))
            ws.cell(row=idx, column=3, value=item.get('content', ''))

            # 设置对齐
            ws.cell(row=idx, column=1).alignment = Alignment(horizontal='center', vertical='top')
            ws.cell(row=idx, column=2).alignment = Alignment(horizontal='center', vertical='top')
            ws.cell(row=idx, column=3).alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    # 冻结首行
    ws.freeze_panes = 'A2'

    # 保存
    wb.save(output_path)
    sp(f"[OK] Excel生成: {output_path.name}")


# ==================== Markdown生成：流派分析 ====================

def generate_approach_word(case_data: Dict, approach_id: str, approach_name: str, output_path: Path):
    """生成流派分析Word文档"""
    doc = Document()

    # 标题
    title = doc.add_heading(f'{approach_name}流派分析', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    case_id = case_data.get('case_id', '')
    case_summary = case_data.get('case_summary', '')

    # 案例基本信息
    doc.add_heading('案例信息', level=1)
    p = doc.add_paragraph()
    p.add_run(f"案例编号：").bold = True
    p.add_run(f"{case_id}\n")
    p.add_run(f"案例概要：").bold = True
    p.add_run(f"{case_summary}")

    # 获取流派分析数据
    analyses = case_data.get('analyses', {})
    approach_data = analyses.get(approach_id, {})

    if not approach_data:
        doc.add_paragraph("暂无该流派的分析数据")
    else:
        # 标签
        tags = approach_data.get('tags', {})
        if tags:
            doc.add_heading('相关标签', level=1)

            relation_tags = tags.get('relation', [])
            if relation_tags:
                p = doc.add_paragraph()
                p.add_run("关系标签：").bold = True
                p.add_run(", ".join(relation_tags))

            symptom_tags = tags.get('symptom', [])
            if symptom_tags:
                p = doc.add_paragraph()
                p.add_run("症状标签：").bold = True
                p.add_run(", ".join(symptom_tags))

        # 危机评估
        crisis_level = approach_data.get('crisis_level', '')
        crisis_evidence = approach_data.get('crisis_evidence', '')
        if crisis_level or crisis_evidence:
            doc.add_heading('危机评估', level=1)
            crisis_labels = {
                'S': 'S级（自杀风险）',
                'A': 'A级（高危）',
                'B': 'B级（中危）',
                'C': 'C级（低危）',
                'D': 'D级（安全）'
            }
            p = doc.add_paragraph()
            p.add_run("等级：").bold = True
            p.add_run(f"{crisis_labels.get(crisis_level, '未知')} ({crisis_level})\n")
            p.add_run("证据：").bold = True
            p.add_run(crisis_evidence)

        # 关键词
        keywords = approach_data.get('keywords', [])
        if keywords:
            doc.add_heading('关键词', level=1)
            doc.add_paragraph(", ".join(keywords))

        # 使用技术
        techniques = approach_data.get('techniques_used', [])
        if techniques:
            doc.add_heading('使用技术', level=1)
            for tech in techniques:
                doc.add_paragraph(tech, style='List Bullet')

        # AI督导分析
        ai_analysis = approach_data.get('ai_analysis', {})
        if ai_analysis:
            doc.add_heading('AI督导分析', level=1)

            summary = ai_analysis.get('summary', '')
            if summary:
                doc.add_heading('概要', level=2)
                doc.add_paragraph(summary)

            strengths = ai_analysis.get('strengths', [])
            if strengths:
                doc.add_heading('优势', level=2)
                for s in strengths:
                    doc.add_paragraph(s, style='List Bullet')

            improvements = ai_analysis.get('improvements', [])
            if improvements:
                doc.add_heading('改进建议', level=2)
                for i in improvements:
                    doc.add_paragraph(i, style='List Bullet')

            recommendations = ai_analysis.get('recommended_followup', [])
            if recommendations:
                doc.add_heading('后续建议', level=2)
                if isinstance(recommendations, list):
                    for r in recommendations:
                        doc.add_paragraph(r, style='List Bullet')
                else:
                    doc.add_paragraph(str(recommendations))

    # 保存
    doc.save(output_path)
    sp(f"[OK] Word生成: {output_path.name}")


def generate_approach_markdown(case_data: Dict, approach_id: str, approach_name: str, output_path: Path):
    """生成流派分析Markdown文档"""

    analyses = case_data.get('analyses', {})
    approach_data = analyses.get(approach_id, {})

    if not approach_data:
        content = f"# {approach_name}分析\n\n暂无分析数据"
    else:
        content = f"# {approach_name}分析\n\n"
        content += f"**案例编号**: {case_data.get('case_id', 'N/A')}\n\n"

        # 危机评估
        crisis_level = approach_data.get('crisis_level', 'Z')
        crisis_evidence = approach_data.get('crisis_evidence', '暂无')
        crisis_labels = {
            "S": "自杀风险", "L": "生命危险", "M": "中度危机",
            "C": "慢性困扰", "Z": "正常范围"
        }
        content += f"## ⚠️ 危机评估\n\n"
        content += f"**等级**: {crisis_labels.get(crisis_level, '未知')} ({crisis_level})\n\n"
        content += f"**证据**: {crisis_evidence}\n\n"

        # 关键词
        keywords = approach_data.get('keywords', [])
        if keywords:
            content += f"## 🏷️ 关键词\n\n"
            content += ", ".join(keywords) + "\n\n"

        # 使用技术
        techniques = approach_data.get('techniques_used', [])
        if techniques:
            content += f"## 🛠️ 使用技术\n\n"
            for tech in techniques:
                content += f"- {tech}\n"
            content += "\n"

        # AI督导分析
        ai_analysis = approach_data.get('ai_analysis', {})
        if ai_analysis:
            content += f"## 🤖 AI督导分析\n\n"

            summary = ai_analysis.get('summary', '')
            if summary:
                content += f"### 概要\n\n{summary}\n\n"

            strengths = ai_analysis.get('strengths', [])
            if strengths:
                content += f"### 优势\n\n"
                for s in strengths:
                    content += f"- {s}\n"
                content += "\n"

            improvements = ai_analysis.get('improvements', [])
            if improvements:
                content += f"### 改进建议\n\n"
                for i in improvements:
                    content += f"- {i}\n"
                content += "\n"

            recommendations = ai_analysis.get('recommendations', [])
            if recommendations:
                content += f"### 后续建议\n\n"
                for r in recommendations:
                    content += f"- {r}\n"
                content += "\n"

    # 保存
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)

    sp(f"[OK] Markdown生成: {output_path.name}")


# ==================== 主函数 ====================

def generate_all_downloads():
    """为所有案例生成下载文件"""

    # 创建输出目录
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # 加载所有案例
    if not CASES_PROCESSED.exists():
        sp("[ERROR] 案例目录不存在")
        return

    case_files = list(CASES_PROCESSED.glob("C*.json"))
    if not case_files:
        sp("[WARN] 没有找到案例文件")
        return

    sp(f"\n开始生成下载文件，共 {len(case_files)} 个案例...\n")

    # 动态加载流派配置
    approaches = []
    config_dir = PROJECT_ROOT / 'data' / 'config' / 'approaches'
    if config_dir.exists():
        for config_file in sorted(config_dir.glob('*.json')):
            try:
                with open(config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                if config.get('enabled', True):
                    approaches.append({
                        "id": config['id'],
                        "name": config.get('name_short') or config['name']
                    })
            except Exception:
                pass
    if not approaches:
        # 降级：默认5个流派
        approaches = [
            {"id": "daguanpai", "name": "大观学派"},
            {"id": "cbt", "name": "CBT"},
            {"id": "psychodynamic", "name": "精神动力学"},
            {"id": "humanistic", "name": "人本主义"},
            {"id": "existential", "name": "存在主义"}
        ]

    for case_file in case_files:
        try:
            with open(case_file, 'r', encoding='utf-8') as f:
                case_data = json.load(f)

            case_id = case_data.get('case_id', 'Unknown')
            sp(f"处理案例: {case_id}")

            # 创建案例专属目录
            case_dir = OUTPUT_DIR / case_id
            case_dir.mkdir(exist_ok=True)

            # 1. 生成复盘Word
            review_path = case_dir / f"{case_id}_复盘.docx"
            generate_review_word(case_data, review_path)

            # 2. 生成逐字稿Excel
            transcript_path = case_dir / f"{case_id}_逐字稿.xlsx"
            generate_transcript_excel(case_data, transcript_path)

            # 3. 生成各流派分析Word
            for approach in approaches:
                word_path = case_dir / f"{case_id}_{approach['name']}.docx"
                generate_approach_word(case_data, approach['id'], approach['name'], word_path)

            sp(f"  ✓ {case_id} 完成\n")

        except Exception as e:
            sp(f"[ERROR] 处理 {case_file.name} 失败: {e}")
            import traceback
            traceback.print_exc()

    sp(f"\n✅ 全部完成！下载文件已生成至: {OUTPUT_DIR}")


if __name__ == "__main__":
    generate_all_downloads()
